# PHES Design App · Snowy 2.0 & Kidston (Sections 5–9)
# Run locally:      streamlit run app.py
# Streamlit Cloud:  push to GitHub → share.streamlit.io → select repo/app.py

import io
import math
import json
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import streamlit as st

# Optional exports
REPORTLAB_OK = True
DOCX_OK = True
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import cm
except Exception:
    REPORTLAB_OK = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:
    DOCX_OK = False

# --------------------------------- Physics helpers ---------------------------------
g = 9.81
rho = 1000.0

def power_MW(Q, h_net, eta):
    return rho * g * Q * h_net * eta / 1e6

def discharge_from_power_MW(P_MW, h_net, eta):
    return P_MW * 1e6 / (rho * g * h_net * eta)

def segment_headloss(L, D, lam, Ksum, Q):
    A = math.pi * D**2 / 4
    v = Q / A if A > 0 else float('nan')
    hf_fric = lam * (L / D) * (v**2 / (2*g)) if D > 0 else float('nan')
    hf_local = Ksum * (v**2 / (2*g))
    return dict(A=A, v=v, hf_fric=hf_fric, hf_local=hf_local, hf=hf_fric + hf_local)

def fit_hf_k_n_from_anchors(h_gross, eta_t, anchors):
    (P1, hf1), (P2, hf2) = anchors
    Q1 = discharge_from_power_MW(P1, h_gross - hf1, eta_t)
    Q2 = discharge_from_power_MW(P2, h_gross - hf2, eta_t)
    n = math.log(hf2 / hf1) / math.log(Q2 / Q1)
    k = hf1 / (Q1 ** n)
    return k, n

def solve_Q_hf_net(P_MW, h_gross, eta_t, k, n, q0=None, tol=1e-10, itmax=200):
    if q0 is None:
        q0 = discharge_from_power_MW(P_MW, h_gross, eta_t)
    Q = q0
    for _ in range(itmax):
        hf = k * (Q ** n)
        hnet = h_gross - hf
        Q_new = discharge_from_power_MW(P_MW, hnet, eta_t)
        if abs(Q_new - Q) < tol:
            Q = Q_new
            break
        Q = 0.5*(Q+Q_new)
    hf = k * (Q ** n)
    return dict(Q=Q, hf=hf, h_net=h_gross - hf)

def hoop_stress_inner(pi_MPa, pext_MPa, ri, re):
    return pi_MPa + 2.0 * pext_MPa * (re**2) / (re**2 - ri**2)

def required_pext_for_ft(pi_MPa, ri, re, ft_MPa):
    return (ft_MPa - pi_MPa) * (re**2 - ri**2) / (2.0 * re**2)

def snowy_vertical_cover(hs, gamma_w=9.81, gamma_R=26.0):
    return (hs * gamma_w) / gamma_R

def norwegian_FRV(CRV, hs, alpha_deg, gamma_w=9.81, gamma_R=26.0):
    return (CRV * gamma_R * math.cos(math.radians(alpha_deg))) / (hs * gamma_w)

def norwegian_FRM(CRM, hs, beta_deg, gamma_w=9.81, gamma_R=26.0):
    return (CRM * gamma_R * math.cos(math.radians(beta_deg))) / (hs * gamma_w)

def surge_tank_first_cut(Ah, Lh, ratio=4.0):
    As = ratio * Ah
    omega_n = math.sqrt(g * Ah / (Lh * As))
    Tn = 2*math.pi/omega_n
    return dict(As=As, omega_n=omega_n, Tn=Tn)

def thoma_sigma_available(H_atm, H_vap, TWL, runner_CL, h_losses_draft, H_net):
    submergence = TWL - runner_CL  # positive if runner below TWL
    return (H_atm - H_vap + submergence - h_losses_draft) / H_net if H_net > 0 else float('nan')

# --------------------------------- UI ---------------------------------
st.set_page_config(page_title="PHES Design (Snowy & Kidston)", layout="wide")
st.title("PHES Design App · Snowy 2.0 & Kidston (Sections 5–9)")

# Presets
with st.sidebar:
    st.header("Presets & Settings")
    preset = st.selectbox("Preset", ["None", "Snowy 2.0 · Plateau (NEW)", "Snowy 2.0 · Plateau (DET)", "Kidston (example)"])
    if st.button("Apply preset"):
        if preset == "Snowy 2.0 · Plateau (NEW)":
            st.session_state.update(dict(HWL_u=1100.0, LWL_u=1080.0, HWL_l=450.0, TWL_l=420.0,
                                         eta_t=0.90, N=6, hf1=28.0, hf2=70.0, P1=1000.0, P2=2000.0))
        elif preset == "Snowy 2.0 · Plateau (DET)":
            st.session_state.update(dict(HWL_u=1100.0, LWL_u=1080.0, HWL_l=450.0, TWL_l=420.0,
                                         eta_t=0.90, N=6, hf1=30.0, hf2=106.0, P1=1000.0, P2=2000.0))
        elif preset == "Kidston (example)":
            st.session_state.update(dict(HWL_u=500.0, LWL_u=490.0, HWL_l=230.0, TWL_l=220.0,
                                         eta_t=0.90, N=2, hf1=6.0, hf2=18.0, P1=250.0, P2=500.0))

    eta_t = st.number_input("Turbine efficiency ηₜ", 0.70, 1.00, float(st.session_state.get("eta_t", 0.90)), 0.01)
    eta_p = st.number_input("Pump efficiency ηₚ (ref.)", 0.60, 1.00, 0.88, 0.01)
    N = st.number_input("Units (N)", 1, 20, int(st.session_state.get("N", 6)), 1)
    st.caption("All units SI; water ρ=1000 kg/m³, g=9.81 m/s².")

# ---------------------------- 5) Levels & rating head ----------------------------
st.subheader("5) Determination of High & Low Water Level")
col1, col2 = st.columns(2)
with col1:
    st.markdown("**Upper reservoir**")
    HWL_u = st.number_input("Upper HWL (m)", 0.0, 3000.0, float(st.session_state.get("HWL_u", 1100.0)), 1.0)
    LWL_u = st.number_input("Upper LWL (m)", 0.0, 3000.0, float(st.session_state.get("LWL_u", 1080.0)), 1.0)
with col2:
    st.markdown("**Lower reservoir**")
    HWL_l = st.number_input("Lower HWL (m)", 0.0, 3000.0, float(st.session_state.get("HWL_l", 450.0)), 1.0)
    TWL_l = st.number_input("Lower TWL (m)", 0.0, 3000.0, float(st.session_state.get("TWL_l", 420.0)), 1.0)

gross_head = HWL_u - TWL_l
NWL_u = HWL_u - (HWL_u - LWL_u)/3.0
head_fluct_rate = (LWL_u - TWL_l)/(HWL_u - TWL_l) if (HWL_u - TWL_l) != 0 else float('nan')

c1, c2, c3 = st.columns(3)
c1.metric("Gross head h_gross (m)", f"{gross_head:.1f}")
c2.metric("NWL (upper) (m)", f"{NWL_u:.1f}")
c3.metric("Head fluctuation rate", f"{head_fluct_rate:.3f}")

with st.expander("Equations"):
    st.latex(r"\text{NWL} \approx \text{HWL} - \frac{\text{HWL}-\text{LWL}}{3}")
    st.latex(r"h_{\text{gross}} = \text{HWL}_{\text{upper}} - \text{TWL}_{\text{lower}}")
    st.latex(r"\text{Head fluctuation rate} = \frac{\text{LWL} - \text{TWL}}{\text{HWL} - \text{TWL}} \ \ (\ge 0.7)")

# ---------------------------- 6) Waterway profile (image OR data) ----------------------------
st.subheader("6) Preparation of Waterway Profile")

left, right = st.columns([2,1])

with right:
    st.markdown("**Option A — Upload a profile image (PNG/JPG)**")
    uploaded_img = st.file_uploader("Profile/figure image", type=["png","jpg","jpeg"], key="profile_img")

with left:
    st.markdown("**Option B — Provide chainage–elevation data**")
    csv_file = st.file_uploader("Upload CSV with columns: Chainage_m, Elevation_m", type=["csv"], key="profile_csv")
    if csv_file is not None:
        df_profile = pd.read_csv(csv_file)
    else:
        # Starter editable table
        _default = pd.DataFrame({
            "Chainage_m": [0, 1000, 2000, 3000, 4000, 5000],
            "Elevation_m": [1085, 1083, 1076, 1040,  980,   920],
        })
        df_profile = st.data_editor(_default, num_rows="dynamic", use_container_width=True, key="profile_editor")

# Plot profile if data present
profile_fig = None
if df_profile is not None and {"Chainage_m","Elevation_m"}.issubset(df_profile.columns):
    ch = df_profile["Chainage_m"].values
    el = df_profile["Elevation_m"].values
    profile_fig = plt.figure(figsize=(7,3))
    plt.plot(ch, el, linewidth=2)
    plt.xlabel("Chainage (m)")
    plt.ylabel("Elevation (m)")
    plt.title("Waterway Long-Section")
    plt.grid(True, linestyle="--", linewidth=0.5)
    st.pyplot(profile_fig)

if uploaded_img is not None:
    st.image(uploaded_img, caption="Uploaded waterway profile image", use_column_width=True)

st.caption("Use either the plotted long-section (data) or the uploaded figure for your report.")

# Runner CL helper / draft head
h_draft = st.number_input("Draft head below lower LWL (m)", 0.0, 200.0, 5.0, 0.5)
runner_CL = LWL_u - h_draft
st.write(f"Runner centreline elevation ≈ **{runner_CL:.1f} m**")

# ---------------------------- 7) Pressure tunnels & shafts ----------------------------
st.subheader("7) Design of Pressure Tunnels & Shafts")

c1, c2, c3, c4 = st.columns(4)
with c1:
    hs = st.number_input("Hydrostatic head to crown h_s (m)", 0.0, 2000.0, 300.0, 1.0)
with c2:
    alpha = st.number_input("Tunnel inclination α (deg)", 0.0, 90.0, 20.0, 1.0)
with c3:
    beta = st.number_input("Ground slope β (deg)", 0.0, 90.0, 40.0, 1.0)
with c4:
    gamma_R = st.number_input("Rock unit weight γ_R (kN/m³)", 15.0, 30.0, 26.0, 0.5)

CRV = snowy_vertical_cover(hs, gamma_w=9.81, gamma_R=gamma_R)
FRV = norwegian_FRV(CRV, hs, alpha, gamma_w=9.81, gamma_R=gamma_R)
CRM = 0.85 * CRV
FRM = norwegian_FRM(CRM, hs, beta, gamma_w=9.81, gamma_R=gamma_R)

c1, c2, c3 = st.columns(3)
c1.metric("C_RV (m)", f"{CRV:.1f}")
c2.metric("F_RV (-)", f"{FRV:.2f}")
c3.metric("F_RM (-)", f"{FRM:.2f}")

with st.expander("Equations"):
    st.latex(r"C_{RV}=\frac{h_s \gamma_w}{\gamma_R},\quad C_{RH}\approx 2C_{RV}")
    st.latex(r"F_{RV}=\frac{C_{RV}\gamma_R \cos\alpha}{h_s \gamma_w} \ge F,\quad F_{RM}=\frac{C_{RM}\gamma_R\cos\beta}{h_s \gamma_w} \ge F")

st.markdown("**Lining stress (thick-wall cylinder)**")
c1, c2, c3, c4 = st.columns(4)
with c1:
    ri = st.number_input("Inner radius r_i (m)", 0.1, 10.0, 3.15, 0.05)
with c2:
    t = st.number_input("Lining thickness t (m)", 0.1, 2.0, 0.35, 0.01)
with c3:
    pi_MPa = st.number_input("Internal pressure p_i (MPa)", 0.0, 10.0, 2.0, 0.1)
with c4:
    pext_manual = st.number_input("External pressure p_ext (MPa)", 0.0, 10.0, 0.0, 0.1)
re = ri + t
ft_MPa = st.number_input("Concrete tensile strength f_t (MPa)", 1.0, 10.0, 3.0, 0.1)

sigma_theta_i = hoop_stress_inner(pi_MPa, pext_manual, ri, re)
pext_req = required_pext_for_ft(pi_MPa, ri, re, ft_MPa)

c1, c2, c3 = st.columns(3)
c1.metric("σ_θ(r_i) (MPa)", f"{sigma_theta_i:.2f}")
c2.metric("p_ext required (MPa)", f"{pext_req:.2f}")
c3.write("Status: " + ("✅ OK (no cracking)" if sigma_theta_i <= ft_MPa else "⚠️ Cracking likely"))

with st.expander("Equations"):
    st.latex(r"\sigma_\theta(r)=A+\frac{B}{r^2},\ \ A=\frac{-p_i r_i^2 + p_{\text{ext}} r_e^2}{r_e^2-r_i^2},\ \ B=\frac{(p_{\text{ext}}+p_i) r_i^2 r_e^2}{r_e^2-r_i^2}")
    st.latex(r"\sigma_\theta(r_i)=p_i+\frac{2 p_{\text{ext}} r_e^2}{r_e^2-r_i^2}")
    st.latex(r"p_{\text{ext,req}}=\frac{(f_t-p_i)(r_e^2-r_i^2)}{2 r_e^2}")

# Cavitation / Thoma σ
st.markdown("**Cavitation (Thoma σ) check**")
col1, col2, col3, col4 = st.columns(4)
with col1:
    H_atm = st.number_input("Atmospheric head H_atm (m)", 0.0, 15.0, 10.33, 0.01)
with col2:
    H_vap = st.number_input("Vapour head H_vap (m)", 0.0, 2.0, 0.30, 0.01)
with col3:
    h_losses_draft = st.number_input("Draft-tube losses (m)", 0.0, 20.0, 1.0, 0.1)
with col4:
    sigma_req = st.number_input("σ_required (vendor)", 0.0, 1.0, 0.10, 0.01)

# ---------------------------- 8) Head losses & effective head ----------------------------
st.subheader("8) Head Loss and Effective Head")

st.markdown("**Anchor-fit (site-level) method**")
colA, colB = st.columns(2)
with colA:
    condition = st.selectbox("Condition", ["Plateau NEW", "Plateau DETERIORATED", "Custom/Kidston"])
    if condition == "Plateau NEW":
        hf1_default, hf2_default = 28.0, 70.0
    elif condition == "Plateau DETERIORATED":
        hf1_default, hf2_default = 30.0, 106.0
    else:
        hf1_default, hf2_default = float(st.session_state.get("hf1", 6.0)), float(st.session_state.get("hf2", 18.0))
with colB:
    P1 = st.number_input("Anchor P1 (MW)", 100.0, 5000.0, float(st.session_state.get("P1", 1000.0)), 10.0)
    P2 = st.number_input("Anchor P2 (MW)", 100.0, 5000.0, float(st.session_state.get("P2", 2000.0)), 10.0)

h_gross_input = st.number_input("Gross head for rating h_gross (m)", 10.0, 3000.0, float(gross_head), 1.0)
hf1 = st.number_input("h_f at P1 (m)", 0.0, 500.0, hf1_default, 0.1)
hf2 = st.number_input("h_f at P2 (m)", 0.0, 500.0, hf2_default, 0.1)

k, n = fit_hf_k_n_from_anchors(h_gross_input, eta_t, [(P1, hf1), (P2, hf2)])
st.write(f"Fitted loss curve:  **h_f = k·Q^n**,  k = {k:.6f},  n = {n:.3f}")

ratings = [2000.0, 2200.0, st.number_input("Custom rating P (MW)", 100.0, 5000.0, 1800.0, 10.0)]
results = []
for P in ratings:
    out = solve_Q_hf_net(P, h_gross_input, eta_t, k, n)
    results.append({"P_MW": P, **out})

df_res = pd.DataFrame(results)
st.dataframe(df_res.style.format({"Q":"{:.2f}", "hf":"{:.2f}", "h_net":"{:.2f}"}), use_container_width=True)

# Cavitation check using first rating
if not df_res.empty:
    H_net_ref = float(df_res.loc[0, "h_net"])
    sigma_av = thoma_sigma_available(H_atm, H_vap, TWL_l, runner_CL, h_losses_draft, H_net_ref)
    st.write(f"Thoma σ_available at {ratings[0]:.0f} MW (H_net={H_net_ref:.1f} m): **{sigma_av:.3f}**")
    st.write("Status: " + ("✅ Meets requirement" if sigma_av >= sigma_req else "⚠️ Below requirement — increase submergence / reduce losses"))

# Plot hf vs Q
fig = plt.figure(figsize=(6,3))
Q_grid = np.linspace(max(1.0, 0.5*df_res["Q"].min()), 1.5*df_res["Q"].max(), 200)
hf_grid = k * Q_grid**n
plt.plot(Q_grid, hf_grid, linewidth=2)
plt.scatter(df_res["Q"], df_res["hf"], marker="x")
plt.xlabel("Q (m³/s)"); plt.ylabel("h_f (m)"); plt.title("Fitted h_f = k·Q^n")
plt.grid(True, linestyle="--", linewidth=0.5)
st.pyplot(fig)

st.markdown("**Segment (Darcy–Weisbach) method – example**")
col1, col2, col3, col4 = st.columns(4)
with col1:
    L_shared = st.number_input("L_shared (m)", 0.0, 50000.0, 157.0, 1.0)
with col2:
    D_shared = st.number_input("D_shared (m)", 0.1, 10.0, 4.8, 0.1)
with col3:
    lam_shared = st.number_input("λ_shared (-)", 0.001, 0.10, 0.015, 0.001, format="%.3f")
with col4:
    Ksum_shared = st.number_input("ΣK_shared (-)", 0.0, 20.0, 4.0, 0.1)

col1, col2, col3, col4 = st.columns(4)
with col1:
    L_pen = st.number_input("L_penstock (m)", 0.0, 50000.0, 106.0, 1.0)
with col2:
    D_pen_in = st.number_input("D_pen_in (m)", 0.1, 10.0, 2.7, 0.1)
with col3:
    D_pen_out = st.number_input("D_pen_out (m)", 0.1, 10.0, 2.1, 0.1)
with col4:
    lam_pen = st.number_input("λ_pen (-)", 0.001, 0.10, 0.018, 0.001, format="%.3f")

# Compute segment losses using first rating
Q_total = float(df_res.loc[0, "Q"]) if not df_res.empty else 0.0
Q_unit = Q_total / N if N > 0 else 0.0
Q_shared = 2 * Q_unit
seg_shared = segment_headloss(L_shared, D_shared, lam_shared, Ksum_shared, Q_shared)
seg_pen = segment_headloss(L_pen, 0.5*(D_pen_in + D_pen_out), lam_pen, 2.0, Q_unit)

df_seg = pd.DataFrame([
    {"segment":"shared (2 units)", **seg_shared},
    {"segment":"penstock (per unit, avg D)", **seg_pen},
])
st.dataframe(df_seg.style.format({"A":"{:.2f}","v":"{:.2f}","hf_fric":"{:.2f}","hf_local":"{:.2f}","hf":"{:.2f}"}), use_container_width=True)

# Velocities
A_shared = math.pi * D_shared**2 / 4
A_in  = math.pi * D_pen_in**2 / 4
A_out = math.pi * D_pen_out**2 / 4
v_shared = (Q_shared / A_shared) if A_shared>0 else float('nan')
v_pen_in = (Q_unit / A_in) if A_in>0 else float('nan')
v_pen_out = (Q_unit / A_out) if A_out>0 else float('nan')

col1, col2, col3 = st.columns(3)
col1.metric("v_shared (m/s)", f"{v_shared:.2f}")
col2.metric("v_pen_in (m/s)", f"{v_pen_in:.2f}")
col3.metric("v_pen_out (m/s)", f"{v_pen_out:.2f}")

# ---------------------------- 9) Surge tanks (first cut) ----------------------------
st.subheader("9) Headrace & Tailrace Surge Tanks (first cut)")
Ah = A_shared
Lh = st.number_input("Headrace length to surge tank L_h (m)", 100.0, 100000.0, 15000.0, 100.0)
ratio = st.number_input("Area ratio A_s/A_h", 1.0, 10.0, 4.0, 0.1)
first_cut = surge_tank_first_cut(Ah, Lh, ratio=ratio)

c1, c2, c3 = st.columns(3)
c1.metric("A_h (m²)", f"{Ah:.1f}")
c2.metric("A_s (m²)", f"{first_cut['As']:.1f}")
c3.metric("Natural period T_n (s)", f"{first_cut['Tn']:.1f}")

with st.expander("Equations"):
    st.latex(r"\frac{A_s}{A_h}\approx 3\text{–}5")
    st.latex(r"\omega_n \approx \sqrt{\frac{g A_h}{L_h A_s}},\quad T_n = \frac{2\pi}{\omega_n}")

# ---------------------------- Downloads ----------------------------
st.markdown("---")
st.subheader("Download results")

# Build bundle
bundle = {
    "inputs": {
        "eta_t": eta_t, "eta_p": eta_p, "N": N,
        "upper": {"HWL": HWL_u, "LWL": LWL_u},
        "lower": {"HWL": HWL_l, "TWL": TWL_l},
        "h_draft": h_draft,
        "lining": {"ri": ri, "t": t, "re": re, "pi_MPa": pi_MPa, "p_ext_manual": pext_manual, "ft_MPa": ft_MPa},
        "anchors": {"condition": condition, "P1": P1, "hf1": hf1, "P2": P2, "hf2": hf2},
        "segments": {"shared": {"L": L_shared, "D": D_shared, "lam": lam_shared, "K": Ksum_shared},
                     "penstock": {"L": L_pen, "D_in": D_pen_in, "D_out": D_pen_out, "lam": lam_pen}},
        "surge": {"Lh": Lh, "ratio": ratio},
        "cavitation": {"H_atm": H_atm, "H_vap": H_vap, "h_losses_draft": h_losses_draft, "sigma_req": sigma_req}
    },
    "derived": {
        "gross_head": gross_head, "NWL_upper": NWL_u, "head_fluct_rate": head_fluct_rate,
        "rock_cover": {"CRV": CRV, "FRV": FRV, "FRM": FRM},
        "lining": {"sigma_theta_ri": sigma_theta_i, "p_ext_required": pext_req},
        "anchor_fit": {"k": k, "n": n, "results": results},
        "segments": df_seg.to_dict(orient="records"),
        "velocities": {"v_shared": v_shared, "v_pen_in": v_pen_in, "v_pen_out": v_pen_out},
        "surge_first_cut": first_cut
    }
}
# Cavitation status
if not df_res.empty:
    H_net_ref = float(df_res.loc[0, "h_net"])
    sigma_av = thoma_sigma_available(H_atm, H_vap, TWL_l, runner_CL, h_losses_draft, H_net_ref)
    bundle["derived"]["cavitation"] = {"sigma_av": sigma_av, "H_net_ref": H_net_ref}

# JSON
st.download_button("Download JSON results",
                   data=json.dumps(bundle, indent=2),
                   file_name="phes_results.json",
                   mime="application/json")

# Excel
excel_bytes = io.BytesIO()
with pd.ExcelWriter(excel_bytes, engine="xlsxwriter") as writer:
    pd.DataFrame([bundle["inputs"]]).to_excel(writer, sheet_name="Inputs", index=False)
    pd.DataFrame(bundle["derived"]["anchor_fit"]["results"]).to_excel(writer, sheet_name="AnchorFit", index=False)
    df_res.to_excel(writer, sheet_name="Ratings", index=False)
    pd.DataFrame(bundle["derived"]["segments"]).to_excel(writer, sheet_name="Segments", index=False)
    pd.DataFrame([bundle["derived"]["lining"]]).to_excel(writer, sheet_name="Lining", index=False)
    pd.DataFrame([bundle["derived"]["rock_cover"]]).to_excel(writer, sheet_name="RockCover", index=False)
    pd.DataFrame([bundle["derived"]["surge_first_cut"]]).to_excel(writer, sheet_name="SurgeFirstCut", index=False)
    if "cavitation" in bundle["derived"]:
        pd.DataFrame([bundle["derived"]["cavitation"]]).to_excel(writer, sheet_name="Cavitation", index=False)
if st.download_button("Download Excel workbook",
                      data=excel_bytes.getvalue(),
                      file_name="phes_results.xlsx",
                      mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"):
    pass

# PDF (ReportLab)
def make_pdf_report(bdl: dict) -> bytes:
    if not REPORTLAB_OK:
        return b''
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    x, y = 2*cm, height - 2*cm
    c.setFont("Helvetica-Bold", 12)
    c.drawString(x, y, "PHES Design Report"); y -= 0.6*cm
    c.setFont("Helvetica", 10)
    lines = [
        f"Gross head h_gross: {bdl['derived']['gross_head']:.1f} m",
        f"NWL upper: {bdl['derived']['NWL_upper']:.1f} m",
        f"Head fluctuation rate: {bdl['derived']['head_fluct_rate']:.3f}",
        f"Rock cover: C_RV={bdl['derived']['rock_cover']['CRV']:.1f} m, F_RV={bdl['derived']['rock_cover']['FRV']:.2f}, F_RM={bdl['derived']['rock_cover']['FRM']:.2f}",
        f"Lining: σθ(ri)={bdl['derived']['lining']['sigma_theta_ri']:.2f} MPa, p_ext,req={bdl['derived']['lining']['p_ext_required']:.2f} MPa",
        f"Anchor-fit: k={bdl['derived']['anchor_fit']['k']:.6f}, n={bdl['derived']['anchor_fit']['n']:.3f}",
    ]
    for ln in lines:
        c.drawString(x, y, ln); y -= 0.5*cm
    c.drawString(x, y, "Ratings:"); y -= 0.5*cm
    for r in bdl["derived"]["anchor_fit"]["results"]:
        c.drawString(x+0.5*cm, y, f"P={r['P_MW']:.0f} MW  Q={r['Q']:.1f} m3/s  hf={r['hf']:.1f} m  h_net={r['h_net']:.1f} m"); y -= 0.5*cm
        if y < 2*cm: c.showPage(); y = height - 2*cm
    c.drawString(x, y, "Segments (first rating):"); y -= 0.5*cm
    for seg in bdl["derived"]["segments"]:
        c.drawString(x+0.5*cm, y, f"{seg['segment']}: v={seg['v']:.2f} m/s, hf={seg['hf']:.2f} m"); y -= 0.5*cm
        if y < 2*cm: c.showPage(); y = height - 2*cm
    if "cavitation" in bdl["derived"]:
        cav = bdl["derived"]["cavitation"]
        c.drawString(x, y, f"Cavitation (Thoma): σ_av={cav['sigma_av']:.3f} @ H_net={cav['H_net_ref']:.1f} m"); y -= 0.5*cm
    c.drawString(x, y, f"Surge first-cut: A_s={bdl['derived']['surge_first_cut']['As']:.1f} m2, T_n={bdl['derived']['surge_first_cut']['Tn']:.1f} s")
    c.showPage(); c.save()
    return buf.getvalue()

if REPORTLAB_OK:
    pdf_bytes = make_pdf_report(bundle)
    st.download_button("Download PDF report", data=pdf_bytes, file_name="phes_report.pdf", mime="application/pdf")
else:
    st.info("Install reportlab to enable PDF export:  pip install reportlab")

# DOCX (python-docx)
def make_docx_report(bdl: dict) -> bytes:
    if not DOCX_OK:
        return b""
    doc = Document()
    doc.add_heading("PHES Design Report", level=1)
    p = doc.add_paragraph()
    p.add_run("Gross head h_gross: ").bold = True; p.add_run(f"{bdl['derived']['gross_head']:.1f} m")
    p = doc.add_paragraph()
    p.add_run("NWL upper: ").bold = True; p.add_run(f"{bdl['derived']['NWL_upper']:.1f} m")
    p = doc.add_paragraph()
    p.add_run("Head fluctuation rate: ").bold = True; p.add_run(f"{bdl['derived']['head_fluct_rate']:.3f}")
    p = doc.add_paragraph()
    p.add_run("Rock cover: ").bold = True
    p.add_run(f"C_RV={bdl['derived']['rock_cover']['CRV']:.1f} m, F_RV={bdl['derived']['rock_cover']['FRV']:.2f}, F_RM={bdl['derived']['rock_cover']['FRM']:.2f}")
    p = doc.add_paragraph()
    p.add_run("Lining: ").bold = True
    p.add_run(f"σθ(ri)={bdl['derived']['lining']['sigma_theta_ri']:.2f} MPa, p_ext,req={bdl['derived']['lining']['p_ext_required']:.2f} MPa")
    p = doc.add_paragraph()
    p.add_run("Anchor-fit: ").bold = True
    p.add_run(f"k={bdl['derived']['anchor_fit']['k']:.6f}, n={bdl['derived']['anchor_fit']['n']:.3f}")

    doc.add_heading("Ratings", level=2)
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'P (MW)'; hdr_cells[1].text = 'Q (m³/s)'; hdr_cells[2].text = 'h_f (m)'; hdr_cells[3].text = 'h_net (m)'
    for r in bdl['derived']['anchor_fit']['results']:
        row_cells = table.add_row().cells
        row_cells[0].text = f"{r['P_MW']:.0f}"
        row_cells[1].text = f"{r['Q']:.1f}"
        row_cells[2].text = f"{r['hf']:.1f}"
        row_cells[3].text = f"{r['h_net']:.1f}"

    doc.add_heading("Segments (first rating)", level=2)
    table2 = doc.add_table(rows=1, cols=3)
    h2 = table2.rows[0].cells
    h2[0].text = 'Segment'; h2[1].text = 'v (m/s)'; h2[2].text = 'h_f (m)'
    for seg in bdl["derived"]["segments"]:
        row = table2.add_row().cells
        row[0].text = seg['segment']
        row[1].text = f"{seg['v']:.2f}"
        row[2].text = f"{seg['hf']:.2f}"

    if "cavitation" in bdl["derived"]:
        doc.add_heading("Cavitation (Thoma)", level=2)
        c = bdl["derived"]["cavitation"]
        doc.add_paragraph(f"σ_available = {c['sigma_av']:.3f} @ H_net={c['H_net_ref']:.1f} m")

    doc.add_heading("Surge first-cut", level=2)
    doc.add_paragraph(f"A_s={bdl['derived']['surge_first_cut']['As']:.1f} m², T_n={bdl['derived']['surge_first_cut']['Tn']:.1f} s")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

if DOCX_OK:
    docx_bytes = make_docx_report(bundle)
    st.download_button("Download Word (DOCX) report", data=docx_bytes, file_name="phes_report.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
else:
    st.info("Install python-docx to enable Word export:  pip install python-docx")

st.caption("Use the profile data (Option B) to auto-plot the long-section if no image is available. Validate final design with transient modelling and vendor data.")
